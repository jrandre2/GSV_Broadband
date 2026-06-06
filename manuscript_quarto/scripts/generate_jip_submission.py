#!/usr/bin/env python3
"""
Generate Journal of Information Policy (JIP) Submission Package.

This script generates all required files for JIP submission:
1. Anonymous article (author info removed)
2. Cover page with author information
3. Separate figure files (300 dpi, TIFF/PNG)
4. Alt text documentation
5. Word count report

Usage:
    python generate_jip_submission.py [--source jip-template.qmd]
"""

import argparse
import re
import shutil
import subprocess
import sys
from datetime import datetime
from pathlib import Path


def count_words_in_docx(docx_path: Path) -> dict:
    """Count words in a DOCX file (requires python-docx)."""
    try:
        from docx import Document
        doc = Document(docx_path)

        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)

        text = '\n'.join(full_text)
        words = len(text.split())

        return {
            'total_words': words,
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
        }
    except ImportError:
        print("Warning: python-docx not installed, skipping word count")
        return {'total_words': 'N/A', 'paragraphs': 'N/A', 'tables': 'N/A'}


def create_anonymous_qmd(source_path: Path, output_path: Path):
    """Create anonymous version of QMD file by removing author information."""
    content = source_path.read_text()

    # Remove author block from YAML header
    # Match author: followed by indented content until next top-level key
    content = re.sub(
        r'^author:.*?(?=^[a-z]|\Z)',
        '',
        content,
        flags=re.MULTILINE | re.DOTALL
    )

    # Remove any "Author Name" placeholders in text
    content = re.sub(r'\[Author Name\]', '[Author]', content)
    content = re.sub(r'Author Name', '[Author]', content)

    # Add anonymous header
    anonymous_header = """# Anonymous Manuscript for Peer Review
# Author information has been removed for blind review.

"""

    # Insert after the YAML front matter
    parts = content.split('---', 2)
    if len(parts) >= 3:
        content = f"---{parts[1]}---\n{anonymous_header}{parts[2]}"

    output_path.write_text(content)
    print(f"Created anonymous QMD: {output_path}")


def export_figures(figures_dir: Path, output_dir: Path):
    """Copy figures to submission folder with proper naming."""
    output_dir.mkdir(parents=True, exist_ok=True)

    figure_files = []
    for ext in ['*.png', '*.jpg', '*.jpeg', '*.tiff', '*.tif']:
        figure_files.extend(figures_dir.glob(ext))

    # Skip macOS metadata and hidden files.
    figure_files = [
        fig_path for fig_path in figure_files
        if not fig_path.name.startswith('._') and not fig_path.name.startswith('.')
    ]

    alt_text_entries = []

    for i, fig_path in enumerate(sorted(figure_files), 1):
        # Create standardized filename
        new_name = f"fig_{i:02d}_{fig_path.stem}{fig_path.suffix}"
        dest_path = output_dir / new_name
        # Copy only file contents to avoid macOS resource forks.
        shutil.copyfile(fig_path, dest_path)
        print(f"Copied figure: {fig_path.name} -> {new_name}")

        # Add placeholder alt text entry
        alt_text_entries.append({
            'filename': new_name,
            'original': fig_path.name,
            'alt_text': f"[REQUIRED: Add alt text description for {fig_path.name}]"
        })

    return alt_text_entries


def create_alt_text_file(alt_text_entries: list, output_path: Path):
    """Create alt text documentation file."""
    content = """# Alt Text for Figures and Tables
# Journal of Information Policy Submission
# =========================================
#
# JIP requires alt text for all figures, charts, and tables.
# Alt text should describe the visual content for accessibility.
#
# Guidelines:
# - Be concise but comprehensive (typically 1-2 sentences)
# - Describe the visual content, not just the title
# - For charts: describe trends, key data points, relationships
# - For maps: describe regions shown, key patterns
# - For tables: summarize structure and key findings

"""

    for entry in alt_text_entries:
        content += f"""
## {entry['filename']}
Original: {entry['original']}
Alt text: {entry['alt_text']}
"""

    content += """

# Tables
# ======
# Add alt text for each table below:

## Table 1: [Table Title]
Alt text: [Describe the table structure and key findings]

## Table 2: [Table Title]
Alt text: [Describe the table structure and key findings]
"""

    output_path.write_text(content)
    print(f"Created alt text file: {output_path}")


def create_word_count_report(word_counts: dict, output_path: Path):
    """Create word count verification report."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    content = f"""# Word Count Report
# Journal of Information Policy Submission
# Generated: {timestamp}
# =========================================

## JIP Requirements
- Maximum: 10,000 words
- Excludes: abstract, tables, references, footnotes

## Document Statistics
- Total words (full document): {word_counts.get('total_words', 'N/A')}
- Paragraphs: {word_counts.get('paragraphs', 'N/A')}
- Tables: {word_counts.get('tables', 'N/A')}

## Compliance Check
"""

    total = word_counts.get('total_words', 0)
    if isinstance(total, int):
        if total <= 10000:
            content += f"[PASS] Word count ({total}) is within limit\n"
        else:
            content += f"[FAIL] Word count ({total}) exceeds 10,000 word limit\n"
            content += f"       Reduce by approximately {total - 10000} words\n"
    else:
        content += "[INFO] Word count could not be verified automatically\n"

    content += """
## Notes
- Word count includes body text only (approximate)
- Manual verification recommended before submission
- Abstract, tables, references, and footnotes are excluded from limit
"""

    output_path.write_text(content)
    print(f"Created word count report: {output_path}")


def render_document(qmd_path: Path, output_dir: Path, manuscript_dir: Path):
    """Render QMD to DOCX using Quarto.

    Because the project uses Quarto manuscript type, we render
    the template in a temporary location outside the project.
    """
    import tempfile

    # Create a temporary directory for rendering
    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)

        # Copy required files to temp directory
        shutil.copy2(qmd_path, temp_path / qmd_path.name)
        shutil.copy2(manuscript_dir / 'jip-reference.docx', temp_path / 'jip-reference.docx')

        # Copy references.bib if it exists
        refs_file = manuscript_dir / 'references.bib'
        if refs_file.exists():
            shutil.copy2(refs_file, temp_path / 'references.bib')

        # Copy CSL directory
        csl_src = manuscript_dir / 'csl'
        csl_dest = temp_path / 'csl'
        if csl_src.exists():
            shutil.copytree(csl_src, csl_dest)

        # Render in temp directory
        temp_qmd = temp_path / qmd_path.name
        cmd = ['quarto', 'render', str(temp_qmd), '--to', 'docx']

        print(f"Rendering: {qmd_path.name}")
        result = subprocess.run(cmd, capture_output=True, text=True, cwd=temp_path)

        if result.returncode != 0:
            print(f"Error rendering document:\n{result.stderr}")
            return None

        # Find and copy the output file (exclude reference.docx)
        docx_files = list(temp_path.glob('*.docx'))
        docx_files = [f for f in docx_files if 'reference' not in f.name.lower()]

        if docx_files:
            output_file = output_dir / docx_files[0].name
            shutil.copy2(docx_files[0], output_file)
            return output_file

    return None


def main():
    parser = argparse.ArgumentParser(
        description='Generate JIP submission package'
    )
    parser.add_argument(
        '--source',
        type=Path,
        default=Path('jip-template.qmd'),
        help='Source QMD file (default: jip-template.qmd)'
    )
    parser.add_argument(
        '--output-dir',
        type=Path,
        default=Path('_output/jip_submission'),
        help='Output directory (default: _output/jip_submission)'
    )
    parser.add_argument(
        '--skip-render',
        action='store_true',
        help='Skip Quarto rendering (use existing DOCX)'
    )

    args = parser.parse_args()

    # Resolve paths relative to manuscript_quarto directory
    script_dir = Path(__file__).parent
    manuscript_dir = script_dir.parent

    source_path = manuscript_dir / args.source
    output_dir = manuscript_dir / args.output_dir
    figures_dir = manuscript_dir / 'figures'

    print(f"\n{'='*60}")
    print("Journal of Information Policy Submission Generator")
    print(f"{'='*60}\n")

    # Create output directory structure
    output_dir.mkdir(parents=True, exist_ok=True)
    (output_dir / 'figures').mkdir(exist_ok=True)

    # Step 1: Create anonymous QMD
    print("\n[Step 1] Creating anonymous manuscript...")
    anon_qmd = output_dir / 'article_anonymous.qmd'
    create_anonymous_qmd(source_path, anon_qmd)

    # Step 2: Render documents
    if not args.skip_render:
        print("\n[Step 2] Rendering documents...")

        # Check if quarto is available
        result = subprocess.run(['quarto', '--version'], capture_output=True)
        if result.returncode != 0:
            print("Error: Quarto not found. Please install Quarto first.")
            print("Visit: https://quarto.org/docs/get-started/")
            sys.exit(1)

        # Render anonymous article
        anon_docx = render_document(anon_qmd, output_dir, manuscript_dir)
        if anon_docx:
            # Rename to standard name
            final_anon = output_dir / 'article_anonymous.docx'
            if anon_docx != final_anon:
                shutil.move(anon_docx, final_anon)
            print(f"Created: {final_anon}")

        # Render cover page
        coverpage_qmd = manuscript_dir / 'jip-coverpage.qmd'
        if coverpage_qmd.exists():
            cover_docx = render_document(coverpage_qmd, output_dir, manuscript_dir)
            if cover_docx:
                final_cover = output_dir / 'cover_page.docx'
                if cover_docx != final_cover:
                    shutil.move(cover_docx, final_cover)
                print(f"Created: {final_cover}")
    else:
        print("\n[Step 2] Skipping render (--skip-render flag set)")

    # Step 3: Export figures
    print("\n[Step 3] Exporting figures...")
    if figures_dir.exists():
        alt_text_entries = export_figures(figures_dir, output_dir / 'figures')
    else:
        print(f"Warning: Figures directory not found: {figures_dir}")
        alt_text_entries = []

    # Step 4: Create alt text file
    print("\n[Step 4] Creating alt text documentation...")
    create_alt_text_file(alt_text_entries, output_dir / 'alt_text.txt')

    # Step 5: Word count report
    print("\n[Step 5] Generating word count report...")
    anon_docx_path = output_dir / 'article_anonymous.docx'
    if anon_docx_path.exists():
        word_counts = count_words_in_docx(anon_docx_path)
    else:
        word_counts = {}
    create_word_count_report(word_counts, output_dir / 'word_count_report.txt')

    # Summary
    print(f"\n{'='*60}")
    print("Submission Package Generated")
    print(f"{'='*60}")
    print(f"\nOutput directory: {output_dir}")
    print("\nContents:")
    for item in sorted(output_dir.rglob('*')):
        if item.is_file():
            rel_path = item.relative_to(output_dir)
            print(f"  - {rel_path}")

    print("\n[NEXT STEPS]")
    print("1. Review article_anonymous.docx for any remaining identifying info")
    print("2. Fill in author details in cover_page.docx")
    print("3. Add alt text descriptions in alt_text.txt")
    print("4. Verify word count is under 10,000")
    print("5. Submit via JIP submission portal")
    print()


if __name__ == '__main__':
    main()
