#!/usr/bin/env python3
"""
Create JIP-compliant Word reference template.

Journal of Information Policy Requirements:
- Times New Roman 12pt body text
- 10pt footnotes
- Double line spacing (2.0)
- 1-inch margins
- Specific heading styles
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def create_jip_reference_template(output_path: Path):
    """Create JIP-compliant Word reference template."""
    doc = Document()

    # Set document margins (1 inch = 914400 EMUs = 1440 twips)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    # Configure Normal style (body text)
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_para = normal_style.paragraph_format
    normal_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal_para.space_after = Pt(0)
    normal_para.space_before = Pt(0)

    # Configure Title style
    title_style = doc.styles['Title']
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(16)
    title_font.bold = True
    title_para = title_style.paragraph_format
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    title_para.space_after = Pt(12)

    # Configure Heading 1 style
    h1_style = doc.styles['Heading 1']
    h1_font = h1_style.font
    h1_font.name = 'Times New Roman'
    h1_font.size = Pt(14)
    h1_font.bold = True
    h1_para = h1_style.paragraph_format
    h1_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h1_para.space_before = Pt(12)
    h1_para.space_after = Pt(6)

    # Configure Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = 'Times New Roman'
    h2_font.size = Pt(12)
    h2_font.bold = True
    h2_font.italic = False
    h2_para = h2_style.paragraph_format
    h2_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h2_para.space_before = Pt(12)
    h2_para.space_after = Pt(6)

    # Configure Heading 3 style
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = 'Times New Roman'
    h3_font.size = Pt(12)
    h3_font.bold = False
    h3_font.italic = True
    h3_para = h3_style.paragraph_format
    h3_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h3_para.space_before = Pt(12)
    h3_para.space_after = Pt(6)

    # Configure Abstract style (if it exists, otherwise create)
    try:
        abstract_style = doc.styles['Abstract']
    except KeyError:
        abstract_style = doc.styles.add_style('Abstract', WD_STYLE_TYPE.PARAGRAPH)
        abstract_style.base_style = doc.styles['Normal']
    abstract_font = abstract_style.font
    abstract_font.name = 'Times New Roman'
    abstract_font.size = Pt(12)
    abstract_para = abstract_style.paragraph_format
    abstract_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    abstract_para.first_line_indent = Inches(0)

    # Configure Block Text style for quotes
    try:
        block_style = doc.styles['Block Text']
    except KeyError:
        block_style = doc.styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
        block_style.base_style = doc.styles['Normal']
    block_font = block_style.font
    block_font.name = 'Times New Roman'
    block_font.size = Pt(11)
    block_para = block_style.paragraph_format
    block_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    block_para.left_indent = Inches(0.5)
    block_para.right_indent = Inches(0.5)

    # Configure Caption style for figures/tables
    try:
        caption_style = doc.styles['Caption']
    except KeyError:
        caption_style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        caption_style.base_style = doc.styles['Normal']
    caption_font = caption_style.font
    caption_font.name = 'Times New Roman'
    caption_font.size = Pt(10)
    caption_para = caption_style.paragraph_format
    caption_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Configure Footnote Text style
    try:
        footnote_style = doc.styles['Footnote Text']
    except KeyError:
        footnote_style = doc.styles.add_style('Footnote Text', WD_STYLE_TYPE.PARAGRAPH)
        footnote_style.base_style = doc.styles['Normal']
    footnote_font = footnote_style.font
    footnote_font.name = 'Times New Roman'
    footnote_font.size = Pt(10)
    footnote_para = footnote_style.paragraph_format
    footnote_para.line_spacing_rule = WD_LINE_SPACING.SINGLE
    footnote_para.space_after = Pt(0)

    # Configure Bibliography style
    try:
        bib_style = doc.styles['Bibliography']
    except KeyError:
        bib_style = doc.styles.add_style('Bibliography', WD_STYLE_TYPE.PARAGRAPH)
        bib_style.base_style = doc.styles['Normal']
    bib_font = bib_style.font
    bib_font.name = 'Times New Roman'
    bib_font.size = Pt(12)
    bib_para = bib_style.paragraph_format
    bib_para.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    bib_para.first_line_indent = Inches(-0.5)
    bib_para.left_indent = Inches(0.5)

    # Add sample content to demonstrate styles
    doc.add_paragraph('Article Title', style='Title')

    doc.add_paragraph('Abstract', style='Heading 1')
    p = doc.add_paragraph(
        'This is a sample abstract paragraph demonstrating the formatting. '
        'The Journal of Information Policy requires a 100-word introductory statement '
        'or a 150-200 word abstract focusing on policy implications.',
        style='Normal'
    )

    doc.add_paragraph('Introduction', style='Heading 1')
    p = doc.add_paragraph(
        'This is body text in Times New Roman 12pt with double line spacing. '
        'The document uses 1-inch margins on all sides.',
        style='Normal'
    )

    doc.add_paragraph('Subsection Example', style='Heading 2')
    p = doc.add_paragraph(
        'This is a level 2 heading example. Headings are bold.',
        style='Normal'
    )

    doc.add_paragraph('Sub-subsection Example', style='Heading 3')
    p = doc.add_paragraph(
        'This is a level 3 heading example. These are italic.',
        style='Normal'
    )

    # Save the document
    doc.save(output_path)
    print(f"Created JIP reference template: {output_path}")


if __name__ == '__main__':
    output_path = Path(__file__).parent.parent / 'jip-reference.docx'
    create_jip_reference_template(output_path)
